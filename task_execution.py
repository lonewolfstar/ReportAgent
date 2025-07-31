import re
import ast
import json
from tqdm import tqdm
from docx import Document
from docx.shared import Inches
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generation_docx(md_text, save_path):
    doc = Document()  # 初始化 Word 文档
    for line in md_text.splitlines():
        line = line.strip()
        if not line:
            continue

        # 标题解析
        if line.startswith('#'):
            doc.add_heading(line[2:], level=1)
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[4:], level=4)

        # 图片标记 <image_X, path>
        elif (m := re.match(r'<image_\d+,\s*(.+?)>', line)):
            img_path = m.group(1).strip()
            try:
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"[图片加载失败：{img_path}]")  # 如果图片不存在，插入提示文字
        # 普通段落
        else:
            doc.add_paragraph(line)

    # 保存
    doc.save(save_path)
    print('Word 已生成：report.docx')


def get_image_index(llm, task, images):
    images_info = ''
    keys = images.keys()
    for key in keys:
        understand = images[key]["understand"]
        images_info = images_info + '\n' + key + ': ' + understand

    # 判断是否需要图片信息
    template = """
            身份定义：你是一个经验丰富的项目报告、学术论文、实验报告撰写者，并且能够有效运用和分配写作相关信息资源。
            章节信息：
                章节标题：{title}
                写作要求：{describe}
                子小节信息：{child}
            图表信息：{images}
            任务需求：需要根据提供的以上信息，判断当前章节的撰写中是否需要查询实验结果图、算法框架图等图表信息，以及该查询哪些图表信息，过滤掉不相关的图表
            特殊限制：
                1. 实验部分撰写必须查询实验结果图
                2. 模型部分撰写必须查询算法框架图
                3. 其余部分自行判断
            返回格式：
                如果需要查询图表，返回相关的需要的图表信息（过滤掉不相关的图表信息），示例：
                ["image_2", "image_3", "image_4", ...]
                如果不需要查询图表，返回： []
            只返回图表名称列表
            """
    prompt = PromptTemplate(template=template, input_variables=['title', 'describe', 'child', 'images'])
    chain = prompt | llm
    response = chain.invoke(
        {'title': task['title'], 'describe': task['describe'], 'child': task['child'], 'images': images_info})
    return ast.literal_eval(response.content)


def get_code_index(llm, task, codes):
    codes_info = ''
    keys = codes.keys()
    for key in keys:
        summarize = codes[key]["summarize"]
        codes_info = codes_info + '\n' + key + ': ' + summarize

    # 判断是否需要代码信息
    template = """
            身份定义：你是一个经验丰富的项目报告、学术论文、实验报告撰写者，并且能够有效运用和分配写作相关信息资源。
            章节信息：
                章节标题：{title}
                写作要求：{describe}
                子小节信息：{child}
            代码信息：{codes}
            任务需求：需要根据提供的以上信息，判断当前章节的撰写中是否需要查询代码信息，以及该查询哪些代码信息，过滤掉不相关的代码信息
            特殊限制：
                1. 实验部分撰写必须查询代码信息
                2. 模型部分撰写必须查询代码信息
                3. 其余部分自行判断
            返回格式：
                如果需要查询代码信息，返回相关的需要的代码信息（过滤掉不相关的代码信息），示例：
                ["code_2", "code_3", "code_4", ...]
                如果不需要查询代码信息，返回： []
            只返回代码信息名称列表
            """
    prompt = PromptTemplate(template=template, input_variables=['title', 'describe', 'child', 'codes'])
    chain = prompt | llm
    response = chain.invoke(
        {'title': task['title'], 'describe': task['describe'], 'child': task['child'], 'codes': codes_info})
    return ast.literal_eval(response.content)


def execution(llm, user, task, images, codes):
    # 获取图片信息
    image_index = get_image_index(llm, task, images)
    images_info = {}
    for index in image_index:
        images_info.update({index: images[index]})

    # 获取代码信息
    code_index = get_code_index(llm, task, codes)
    codes_info = {}
    for index in code_index:
        codes_info.update({index: codes[index]['describe']})

    template = """
            身份定义：你是一个经验丰富的项目报告、学术论文、实验报告撰写者，擅长根据提供的信息创作高质量内容报告。
            章节信息：
                章节标题：{title}
                写作要求：{describe}
                子小节信息：{child}
            图表信息：{images}
            代码信息：{codes}
            全局信息（只参考与当前章节相关部分）：
                项目背景：{project}
                数据描述：{dataset}
                写作风格：{style}
            任务需求：
                你需要根据提供的以上信息创建一个完整的内容报告。报告应该遵循以下几点:
                    1. 保持清晰的段落结构和逻辑流程，如果有子小节，按子小节格式撰写
                    2. 使用指定的写作风格
                    3. 内容应当丰富、连贯且有价值
                    4. 对于全局信息的参考，只参考与当前章节相关部分
                创建的报告应该像一个完整的成品，而不仅仅是填充大纲。

                (判断)在撰写过程中如果有图表信息参考，需要融入图表信息并且在合理的位置上插入图像，使用<图像名称, 图像路径>标记表示图片插入的位置，示例：
                    ......
                    下图展示了训练集上BERT模型的混淆矩阵。分析结果为：消极评论（0）识别效果良好，有53091个样本被正确识别，但有6880个消极评论样本被误分为积极评论（1）；积极评论（1）识别效果良好，有52805个样本被正确识别，但有7160个积极评论样本被误分为消极评论（0）。 
                        <image_9, ./upload_files/analysis/train/bert/confusion_matrix.png>
                    混淆矩阵表明，在验证集和外部队列1中，模型对正常和良性疾病的预测较为准确，但存在恶性肿瘤误判为良性疾病的倾向。总体而言，模型在区分正常类别方面表现出色，并且能够准确地识别大部分良性疾病和恶性肿瘤，但仍存在良性疾病与恶性肿瘤区分错误的情况。
                    ......
                如果没有图表信息参考，不用考虑。
                
                (判断)在撰写过程中如果有代码信息参考，需要融入代码信息并且强调代码信息中的参数数值，示例：
                    ......
                    在模型训练策略方面，多尺度特征融合模型采用RAdam优化器，初始学习率设为1e-3，并通过指数学习率调度器进行动态调整，批量大小为256，训练轮次为30。多视角特征融合模型和多模态特征融合模型则均使用SGD优化器，初始学习率同样为1e-3，并借助指数学习率调度器实现动态调整。其中，多视角特征融合模型的批量大小为512，训练轮次为50；多模态特征融合模型的批量大小为180，训练轮次为30。
                    ......
                如果没有代码信息参考，不用考虑。
            """
    prompt = PromptTemplate(template=template,
                            input_variables=['project', 'dataset', 'style', 'title', 'describe', 'child',
                                             'images', 'codes'])
    chain = prompt | llm
    response = chain.invoke({'project': user['project'],
                             'dataset': user['dataset'],
                             'style': user['style'],
                             'title': task['title'],
                             'describe': task['describe'],
                             'child': task['child'],
                             'images': images_info,
                             'codes': codes_info})
    return response.content


if __name__ == "__main__":
    # 加载大模型
    llm = ChatOpenAI(
        model='qwen-plus',
        openai_api_base='https://dashscope.aliyuncs.com/compatible-mode/v1',
        openai_api_key='sk-d08ca9f46d8f4a93bf0cafbb9b221ca8',
        temperature=0.3
    )

    # 用户输入描述
    user_input = {
        'project': '构建了Bert模型和Bert_LSTM模型进行商品评论情感分类，对用户的商品评论进行分类，分为消极（负类）和积极（正类）',
        'dataset': '本研究采用的多源数据集整合京东用户行为数据与微博社交媒体的公开文本，构建了一个包含超过10万+条文本样本的综合数据集。数据来源分为以下两个部分：一方面，从微博平台抓取了公开讨论内容，形成了名为weibo_senti_100k.csv的数据集。这部分数据包含了实时话题讨论和公众观点，能够反映社交媒体上的即时情感动态。另一方面通过京东商品评论，形成了名为jd_conm.csv的数据集。最后通过分层采样处理了主流新闻媒体和网络论坛中的时事评论与热点讨论，构成了两个数据集：train.csv（weibo_senti_100k）和test.csv（jd_conm清洗后数据）。',
        'requirements': '1. 背景部分,经典三段式，连贯不要随意插入小标题并且字数不少于800字； 2. 模型描述部分，主要描述模型的算法原理（尽量包含相应的公式，latex格式表示），不要出现具体代码，整体字数不少于800字',
        'style': '学术论文'
    }

    # 加载任务信息
    with open('./task.json', 'r', encoding='utf-8') as f:
        task = json.load(f)

    # 加载图信息
    with open('./images.json', 'r', encoding='utf-8') as f:
        images = json.load(f)

    # 加载代码信息
    with open('./codes.json', 'r', encoding='utf-8') as f:
        codes = json.load(f)

    all_content = ''
    keys = task.keys()
    for key in tqdm(keys, desc="Processing tasks"):
        t = task[key]
        content = execution(llm, user_input, t, images, codes)
        all_content += '\n' + content
    generation_docx(all_content, './report.docx')
